"""Build Word (.docx) insight reports from structured synthesis output (pure Python)."""

from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from pharma_rd.config import Settings
from pharma_rd.integrations.report_executive_framing import prepare_run_summary_for_report
from pharma_rd.pipeline.contracts import SynthesisOutput


def _add_plain_para(doc: Document, text: str, *, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(11)


def _fr26_plain(settings: Settings) -> str | None:
    if settings.deployment_profile != "practice":
        return None
    return (
        "Deployment (FR26): Practice build — public/mock sources permitted; "
        "enterprise SSO is not required for this MVP (FR34 is roadmap)."
    )


def build_insight_report_docx(
    run_id: str,
    synthesis: SynthesisOutput,
    settings: Settings,
) -> bytes:
    """Return ``.docx`` bytes aligned with the Markdown/HTML template report."""
    org = (settings.insight_org_display_name or "").strip() or "Organization"
    doc = Document()
    t = doc.add_heading(f"Insight report — {org}", level=0)
    t.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    sub = doc.add_paragraph()
    sub.add_run(f"Run ID: {run_id}").bold = True

    doc.add_heading("Run summary", level=1)
    _add_plain_para(
        doc,
        "Human judgment (FR22): Items in this report are recommendations for "
        "review only—not approvals. Pursuit and portfolio decisions remain "
        "human-owned.",
    )
    fr26 = _fr26_plain(settings)
    if fr26:
        _add_plain_para(doc, fr26)

    doc.add_paragraph().add_run("Signal characterization: ").bold = True
    doc.add_paragraph(synthesis.signal_characterization)

    rsum = prepare_run_summary_for_report(synthesis)
    doc.add_paragraph().add_run("Monitoring snapshot (FR28)").bold = True
    if rsum.humanized_scan_lines:
        for line in rsum.humanized_scan_lines:
            doc.add_paragraph(line, style="List Bullet")
    else:
        doc.add_paragraph("(No scan summary lines — legacy or empty synthesis.)")

    if rsum.coverage_gap_lines:
        doc.add_paragraph().add_run(
            "Coverage and configuration gaps (preview)"
        ).bold = True
        for g in rsum.coverage_gap_lines:
            doc.add_paragraph(g, style="List Bullet")
        if rsum.coverage_remaining:
            doc.add_paragraph(f"… ({rsum.coverage_remaining} more)")
    if rsum.operator_gap_lines:
        doc.add_paragraph().add_run(
            "Technical pipeline notes (operators)"
        ).bold = True
        _add_plain_para(
            doc,
            "Internal telemetry for this run; JSON artifacts hold full detail.",
            italic=True,
        )
        for g in rsum.operator_gap_lines:
            doc.add_paragraph(g, style="List Bullet")
        if rsum.operator_remaining:
            doc.add_paragraph(f"… ({rsum.operator_remaining} more)")

    doc.add_heading("Ranked opportunities", level=1)
    if not synthesis.ranked_opportunities:
        _add_plain_para(doc, "(none)")
    else:
        for row in sorted(synthesis.ranked_opportunities, key=lambda r: r.rank):
            doc.add_heading(f"{row.rank}. {row.title}", level=2)
            _add_plain_para(
                doc,
                "Recommendation only—not an approval. Pursuit is a human decision.",
                italic=True,
            )
            doc.add_paragraph().add_run("Rationale").bold = True
            _add_plain_para(doc, row.rationale_short)
            if row.evidence_references:
                doc.add_paragraph().add_run("Evidence references").bold = True
                for er in row.evidence_references:
                    doc.add_paragraph(
                        f"{er.domain} — {er.label}: {er.reference}",
                        style="List Bullet",
                    )
            doc.add_paragraph().add_run("Commercial viability").bold = True
            _add_plain_para(doc, row.commercial_viability or "—")

    doc.add_heading("Governance and disclaimer", level=1)
    _add_plain_para(
        doc,
        "This report presents recommendations derived from signals available to the "
        "workflow. It does not approve development, launch, or commercialization. "
        "Pursuit decisions remain with qualified human decision-makers and your "
        "organization's governance processes. Use this as input to judgment, not a "
        "substitute for it.",
    )
    _add_plain_para(
        doc,
        "Recommendations, not approvals—pursuit decisions remain human-owned (FR22).",
        italic=True,
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
